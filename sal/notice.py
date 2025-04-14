import streamlit as st
from datetime import datetime, timedelta
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from PIL import Image
import tempfile
from docx.shared import Cm

# Sidebar Inputs
st.sidebar.header("Inputs")

branch = st.sidebar.selectbox("Branch", ["ME", "ME Minor CSE", "AE"])
department = "Mechanical Engineering"

semester_ref = st.sidebar.selectbox("Semester Ref", ["S1", "S2", "S3", "S4", "S5", "S6", "S7"])
semester = st.sidebar.selectbox("Semester", list(range(1, 8)))

batch = st.sidebar.selectbox("Batch", list(range(2020, 2036)))
year = st.sidebar.selectbox("Year", list(range(2020, 2036)), help="Year of Semester")

session_options = [f"Jul-Dec-{y}" for y in range(2024, 2034)] + [f"Jan-Jun-{y}" for y in range(2025, 2035)]
session = st.sidebar.selectbox("Session", sorted(session_options))

# Dates
date_st1 = st.sidebar.date_input("Date ST1", help="Date after ST1")
date_st2 = st.sidebar.date_input("Date ST2", help="Date after ST2")
date_after_st2 = st.sidebar.date_input("Date after ST2 Result", help="Date after ST2 Result Declaration")
date_after_ete = st.sidebar.date_input("Date after ETE Result", help="Date after ETE Result Declaration")

# Derived dates
date_st1_plus1 = date_st1 + timedelta(days=1)
date_st2_plus1 = date_st2 + timedelta(days=1)

# Thresholds
slow_threshold = st.sidebar.number_input("Slow Threshold %", min_value=40.0, value=40.0)
advanced_threshold = st.sidebar.number_input("Advanced Threshold %", min_value=50.0, value=75.0)

# Subjects
st.sidebar.markdown("**List of Subjects ST1**")
subjects_st1 = [st.sidebar.text_input(f"Subject {i+1} (ST1)", "") for i in range(10)]
subjects_st1_list = "\n".join([f"- {subj}" for subj in subjects_st1 if subj.strip() != ""])

st.sidebar.markdown("**List of Subjects ST2**")
subjects_st2 = [st.sidebar.text_input(f"Subject {i+1} (ST2)", "") for i in range(10)]
subjects_st2_list = "\n".join([f"- {subj}" for subj in subjects_st2 if subj.strip() != ""])

# Header Image
image_file = st.sidebar.file_uploader("Upload Header Image (Top Right)", type=["png", "jpg", "jpeg"])

def add_header_and_style(doc):
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    if image_file:
        for section in doc.sections:
            header = section.header

            # Total width for the table (approx A4 width minus margins)
            table = header.add_table(rows=1, cols=2, width=Cm(17))
            table.allow_autofit = True

            # Set fixed column widths
            table.columns[0].width = Cm(14)  # Push image to the right
            table.columns[1].width = Cm(3)

            cell_left, cell_right = table.rows[0].cells
            cell_left.text = ""  # Leave left cell empty

            # Add image to the right cell
            run = cell_right.paragraphs[0].add_run()
            img = Image.open(image_file)
            with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as temp_img:
                img.save(temp_img.name)
                run.add_picture(temp_img.name, width=Cm(2.5))

            # Align image to the right edge
            cell_right.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

def add_signature(doc):
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "\n\nDean"
    table.rows[0].cells[1].text = "\n\nMentor"
    table.rows[0].cells[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

def add_notice(doc, ref, date, subject, content):
    # para = doc.add_paragraph()
    # run = para.add_run(f"Ref. No.: CUIET/MED/SAL/{year}/{semester_ref}/{branch}/{ref}Date: {date.strftime('%d-%m-%Y')}")
    # run.bold = True
    # para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    
    # Add a table with 1 row and 2 columns
    table = doc.add_table(rows=1, cols=2)

    # Access the cells of the table
    cell_ref = table.cell(0, 0)
    cell_date = table.cell(0, 1)

    # Add the reference text in the left cell (cell_ref) and make it bold
    run_ref = cell_ref.paragraphs[0].add_run(f"Ref. No.: CUIET/MED/SAL/{year}/{semester_ref}/{branch}/{ref}")
    run_ref.bold = True

    # Add the date text in the right cell (cell_date) and make it bold
    run_date = cell_date.paragraphs[0].add_run(f"Date: {date.strftime('%d-%m-%Y')}")
    run_date.bold = True

    # Optionally, align the text to the left and right
    cell_ref.paragraphs[0].alignment = 0  # Left-aligned
    cell_date.paragraphs[0].alignment = 2  # Right-aligned
    
    # Add bold and center-aligned paragraphs
    para1 = doc.add_paragraph(f"\nDepartment of {department}")
    para1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    para1.runs[0].bold = True  # Make the text bold

    para2 = doc.add_paragraph("CUIET – Applied Engineering")
    para2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    para2.runs[0].bold = True  # Make the text bold

    para3 = doc.add_paragraph("\nNotice")
    para3.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    para3.runs[0].bold = True  # Make the text bold

    para4 = doc.add_paragraph(f"\nSubject: {subject}\n")
    para4.runs[0].bold = True  # Make the text bold

    para5 = doc.add_paragraph(content)
    # You can apply bold to the content if needed, but here it's not bold by default.

    para6 = doc.add_paragraph("\nNote: Mentors are requested to inform the above students.")
    para6.runs[0].bold = True  # Make the text bold

    # Assuming add_signature function is defined elsewhere
    add_signature(doc)

    para7 = doc.add_paragraph("\ncc:\n-Notice Board\n-Departmental File\n-Mentoring File")
    para7.runs[0].bold = False  # Make the text bold

    doc.add_page_break()

def generate_doc():
    doc = Document()
    
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)
    
    add_header_and_style(doc)

    add_notice(
        doc, "01", date_st1,
        "Identification of Slow and Advanced Learners",
        f"The students of {branch} {session} were classified into the Slow and Advanced learners' categories based upon the observations and feedback from the mentors, teachers and academic performance in ST-I. Students, who score marks below {slow_threshold}% are categorized as slow learners and above {advanced_threshold}% are categorized as advanced learners. These distinguished parameters enabled in identification of advanced learners and slow learners. The details of slow and advanced learners is available in Annexure A1."
    )

    add_notice(
        doc, "02", date_st1,
        "Schedule of extra classes",
        f"List of Subjects to be offered:\n{subjects_st1_list}\n\nNote: The extra classes on all non-teaching working days are being offered to the semester {semester} {branch} students (Slow Learners). Details of the extra classes are available in Annexure B1."
    )

    para1 = doc.add_paragraph(f"\nDepartment of {department}")
    para1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    para1.runs[0].bold = True  # Make the text bold

    para2 = doc.add_paragraph("CUIET – Applied Engineering")
    para2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    para2.runs[0].bold = True  # Make the text bold

    para3 = doc.add_paragraph("\nAction Taken Report")
    para3.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    para3.runs[0].bold = True  # Make the text bold

    doc.add_paragraph(
        f"A Departmental Meeting was held to discuss the provision to be made and to formulate the adapted teaching methodology for the slow and advanced learners of semester {semester} (Batch: {batch}) students\n\nAction Taken for Slow Learners:\nExtra Classes\nList of Subjects to be offered:\n{subjects_st1_list}\n\nNote: The Classes will be offered as per the Schedule given for various subjects in reference no. CUIET/MED/SAL/{year}/{semester_ref}/{branch}/02.\n\nTeaching Methodology:\nThe lectures were supported with advanced visual modes of teaching and learning in order to boost learning capabilities of students. Extra initiatives were taken up to regulate their performance metrics through regular assignment, viva-voce and quiz sessions etc."
    )
    doc.add_paragraph("\nAction taken for Advanced Learners:\nThe students who have attained advanced learner level were motivated to learn the advanced courses through NPTEL Lectures (links were provided to them) to enrich their skills.")
    # Add a table with 1 row and 2 columns (for two columns layout)
    table = doc.add_table(rows=1, cols=2)

    # Access the cells of the table
    cell_left = table.cell(0, 0)
    cell_right = table.cell(0, 1)

    # Add the left column text (Date and Program Incharge)
    para_left = cell_left.paragraphs[0]
    para_left.add_run(f"Date: {date_st1_plus1.strftime('%d-%m-%Y')}\nProgram Incharge")
    para_left.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # Left-justified

    # Add the right column text (Program Incharge and Department)
    para_right = cell_right.paragraphs[0]
    para_right.add_run(f"Program Incharge\nDepartment of {department}")
    para_right.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT  # Right-justified
    doc.add_page_break()

    para1 = doc.add_paragraph(f"\nDepartment of {department}")
    para1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    para1.runs[0].bold = True  # Make the text bold

    para2 = doc.add_paragraph("CUIET – Applied Engineering")
    para2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    para2.runs[0].bold = True  # Make the text bold

    para3 = doc.add_paragraph("\nPerformance Report")
    para3.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    para3.runs[0].bold = True  # Make the text bold
    doc.add_paragraph(
        f"Slow Learners:\nThe Extra classes were provided to the students who had been identified as slow learners (on the basis of marks obtained in ST-1). After attending these classes a considerable improvement of major number of students is reflected in the grades obtained by them in ST-2. Details of slow learners who showed improvement in ST-2 are available in Annexure D1.\n\nAdvanced Learners: The students who have attained advanced learner level were motivated to learn the advanced courses through NPTEL Lectures to enrich their skills."
    )
    # Add the left column text (Date and Program Incharge)
    para_left = cell_left.paragraphs[0]
    para_left.add_run(f"Date: {date_st1_plus1.strftime('%d-%m-%Y')}\nProgram Incharge")
    para_left.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # Left-justified

    # Add the right column text (Program Incharge and Department)
    para_right = cell_right.paragraphs[0]
    para_right.add_run(f"Program Incharge\nDepartment of {department}")
    para_right.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT  # Right-justified
    doc.add_page_break()

    add_notice(
        doc, "03", date_st2,
        "Identification of Slow and Advanced Learners",
        f"The students of {branch} {session} were classified into the Slow and Advanced learners' categories based upon the observations and feedback from the mentors, teachers and academic performance in ST-I and ST-II. Students, who score marks below {slow_threshold}% in ST-2 are categorized as slow learners and above {advanced_threshold}% in ST-2 are categorized as advanced learners. These distinguished parameters enabled in identification of advanced learners and slow learners. The details of slow and advanced learners is available in Annexure A2."
    )

    add_notice(
        doc, "04", date_st2,
        "Schedule of extra classes",
        f"List of Subjects to be offered:\n{subjects_st2_list}\n\nNote: The extra classes are being offered to the semester {semester} {branch} students (Slow Learners). Details of the extra classes are available in Annexure B2."
    )

    para1 = doc.add_paragraph(f"\nDepartment of {department}")
    para1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    para1.runs[0].bold = True  # Make the text bold

    para2 = doc.add_paragraph("CUIET – Applied Engineering")
    para2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    para2.runs[0].bold = True  # Make the text bold

    para3 = doc.add_paragraph("\nAction Taken Report")
    para3.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    para3.runs[0].bold = True  # Make the text bold
    doc.add_paragraph(
        f"A Departmental Meeting was held to discuss the provision to be made and to formulate the adapted teaching methodology for the slow and advanced learners of semester {semester} (Batch: {batch}) students\n\nAction Taken for Slow Learners:\nExtra Classes\nList of Subjects to be offered:\n{subjects_st2_list}\n\nNote: The Classes will be offered as per the Schedule given for various subjects in reference no. CUIET/MED/SAL/{year}/{semester_ref}/{branch}/04.\n\nTeaching Methodology:\nLectures supported with advanced visual modes, regular assignments and quizzes, previous year papers and question banks provided."
    )
    doc.add_paragraph("\nAction taken for Advanced Learners:\nMOOC courses, expert talks, good project participation, technical club membership.")
    # Add the left column text (Date and Program Incharge)
    para_left = cell_left.paragraphs[0]
    para_left.add_run(f"Date: {date_st1_plus1.strftime('%d-%m-%Y')}\nProgram Incharge")
    para_left.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # Left-justified

    # Add the right column text (Program Incharge and Department)
    para_right = cell_right.paragraphs[0]
    para_right.add_run(f"Program Incharge\nDepartment of {department}")
    para_right.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT  # Right-justified
    doc.add_page_break()

    para1 = doc.add_paragraph(f"\nDepartment of {department}")
    para1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    para1.runs[0].bold = True  # Make the text bold

    para2 = doc.add_paragraph("CUIET – Applied Engineering")
    para2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    para2.runs[0].bold = True  # Make the text bold

    para3 = doc.add_paragraph("\nPerformance Report")
    para3.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    para3.runs[0].bold = True  # Make the text bold
    doc.add_paragraph(
        f"Slow Learners:\nExtra classes were provided to students identified as slow learners (based on ST-1 & ST-II). A considerable improvement was noted in end term results. Details are available in Annexure D2.\n\nAdvanced Learners: Motivated to pursue NPTEL and similar advanced learning paths."
    )
    # Add the left column text (Date and Program Incharge)
    para_left = cell_left.paragraphs[0]
    para_left.add_run(f"Date: {date_st1_plus1.strftime('%d-%m-%Y')}\nProgram Incharge")
    para_left.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # Left-justified

    # Add the right column text (Program Incharge and Department)
    para_right = cell_right.paragraphs[0]
    para_right.add_run(f"Program Incharge\nDepartment of {department}")
    para_right.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT  # Right-justified

    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(temp_file.name)
    return temp_file.name

# Generate button
if st.button("Generate Document"):
    output_path = generate_doc()
    with open(output_path, "rb") as f:
        st.download_button("Download Word Document", f, file_name=f"{semester}.docx")