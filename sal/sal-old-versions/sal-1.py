import streamlit as st
import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from io import BytesIO
import zipfile

# ========== SIDEBAR SETTINGS ==========
st.sidebar.header("📊 Learner Classification Settings")

slow_thresh = st.sidebar.number_input("Slow Learner Threshold (%)", min_value=0, max_value=100, value=40)
advanced_thresh = st.sidebar.number_input("Advanced Learner Threshold (%)", min_value=0, max_value=100, value=75)

uploaded_files = st.sidebar.file_uploader("📂 Upload Excel files", type=["xlsx"], accept_multiple_files=True)

# ========== FUNCTIONS ==========

def classify_student(percentage):
    if percentage < slow_thresh:
        return 'Slow Learner'
    elif percentage > advanced_thresh:
        return 'Advanced Learner'
    else:
        return 'Average Learner'

def sanitize_percent(value):
    str_val = str(value).strip().lower()
    if str_val in ['a', 'absent']:
        return 0.0
    try:
        return float(value)
    except:
        return 0.0

def process_file(uploaded_file):
    df = pd.read_excel(uploaded_file)

    required_columns = [
        'Roll No.', 'Student Name', 'Subject Code', 'Subject Name',
        'Branch', 'Batch', 'ST1 Percentage', 'ST2 Percentage'
    ]
    df = df[required_columns]

    df['ST1 Percentage'] = df['ST1 Percentage'].apply(sanitize_percent).astype(float)
    df['ST2 Percentage'] = df['ST2 Percentage'].apply(sanitize_percent).astype(float)

    df['ST1 Status'] = df['ST1 Percentage'].apply(classify_student)
    df['ST2 Status'] = df['ST2 Percentage'].apply(classify_student)

    return df

def set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.append(tblPr)

    tblBorders = OxmlElement('w:tblBorders')

    for edge in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        elem = OxmlElement(f'w:{edge}')
        elem.set(qn('w:val'), 'single')
        elem.set(qn('w:sz'), '4')
        elem.set(qn('w:space'), '0')
        elem.set(qn('w:color'), '000000')
        tblBorders.append(elem)

    tblPr.append(tblBorders)

def style_table_font(cell, bold=False):
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run.font.color.rgb = None
            run.bold = bold

def add_table_to_doc(doc, group, test_col, status_col, test_name):
    title = f'{test_name} Report'
    title_para = doc.add_paragraph(title)
    run = title_para.runs[0]
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

    table = doc.add_table(rows=1, cols=8)
    set_table_borders(table)

    headers = ['Roll No.', 'Student Name', 'Subject Code', 'Subject Name',
               'Branch', f'{test_name} Percentage', 'Status', 'Batch']
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        style_table_font(hdr_cells[i], bold=True)

    for _, row in group.iterrows():
        row_cells = table.add_row().cells
        values = [
            row['Roll No.'], row['Student Name'], row['Subject Code'], row['Subject Name'],
            row['Branch'], row[test_col], row[status_col], row['Batch']
        ]
        for i, val in enumerate(values):
            if isinstance(val, float):
                row_cells[i].text = f"{val:.1f}"
            else:
                row_cells[i].text = str(val)
            style_table_font(row_cells[i])

    doc.add_paragraph()  # spacing

def generate_grouped_docs(df):
    zip_buffer = BytesIO()
    with zipfile.ZipFile(zip_buffer, "w") as zipf:
        grouped = df.groupby(['Batch', 'Branch'])
        for (batch, branch), group in grouped:
            doc = Document()
            doc.styles['Normal'].font.name = 'Times New Roman'
            doc.styles['Normal'].font.size = Pt(12)

            add_table_to_doc(doc, group, 'ST1 Percentage', 'ST1 Status', 'ST1')
            add_table_to_doc(doc, group, 'ST2 Percentage', 'ST2 Status', 'ST2')

            doc_buffer = BytesIO()
            doc.save(doc_buffer)
            doc_buffer.seek(0)

            filename = f"{batch}_{branch}_Learner_Report.docx"
            zipf.writestr(filename, doc_buffer.read())

    zip_buffer.seek(0)
    return zip_buffer

# ========== MAIN DISPLAY AREA ==========

st.title("📚 Grouped Student Reports by Batch & Branch")

if uploaded_files:
    all_data = pd.DataFrame()

    for file in uploaded_files:
        df = process_file(file)
        all_data = pd.concat([all_data, df], ignore_index=True)

    st.success("✅ Files processed successfully!")

    final_zip = generate_grouped_docs(all_data)

    st.download_button(
        label="📦 Download All Reports as ZIP",
        data=final_zip,
        file_name="All_Learner_Reports.zip",
        mime="application/zip"
    )
