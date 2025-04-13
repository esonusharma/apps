import streamlit as st
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from io import BytesIO

def generate_placeholder_doc(data, image_file=None):
    doc = Document()

    # Set font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    # Page 1
    if image_file:
        table_img = doc.add_table(rows=1, cols=2)
        img_cells = table_img.rows[0].cells
        img_cells[0].text = ""
        run = img_cells[1].paragraphs[0].add_run()
        run.add_picture(image_file, width=Inches(1.0))
        img_cells[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    table = doc.add_table(rows=1, cols=2)
    cells = table.rows[0].cells
    cells[0].text = f"Ref. No.: CUIET/MED/SAL/{data['year']}/{data['semester']}/<<notice11>>"
    cells[1].text = f"Date: {data['date']}"
    cells[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    doc.add_paragraph(f"\nDepartment of <<department>>").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph("CUIET – Applied Engineering").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph("\nNotice").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph("\nSubject: Identification of Slow and Advanced Learners\n")
    doc.add_paragraph(
        "The below mentioned students of <<Branch>> <<session>> session July –Dec 2023 were classified "
        "into the Slow and Advanced learners' categories based upon the observations and feedback from "
        "the mentors, teachers and academic performance in <<st>> ST-I. Students, who score marks below "
        "<<slow threshold>> are categorized as slow learners and above <<advanced threshold>> are categorized "
        "as advanced learners. The details of slow and advanced learners is available in Annexure A1."
    )
    doc.add_paragraph("\nNote: Mentors are requested to inform the above students.")
    doc.add_table(rows=1, cols=2).rows[0].cells[0].text = "\n\nDean"
    doc.tables[-1].rows[0].cells[1].text = "\n\nMentor"
    doc.tables[-1].rows[0].cells[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    doc.add_paragraph("\ncc:\n-Notice Board\n-Departmental File\n-Mentoring File")

    doc.add_page_break()

    # Page 2
    doc.add_paragraph(f"Ref. No.: CUIET/MED/SAL/{data['year']}/{data['semester']}/<<notice12>>\t\t\t\tDate: {data['date']}")
    doc.add_paragraph(f"\nDepartment of <<department>>").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph("CUIET – Applied Engineering").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph("\nNotice").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph("\nSubject: Schedule of extra classes\n\nList of Subjects to be offered\n\n<<list of subjects st1>>")
    doc.add_paragraph("Note: The extra classes on all non-teaching working days are being offered to the 3rd semester <<branch>> students (Slow Learners).")
    doc.add_paragraph("\n\nSubject\nSemester\nDate\nTime\nVenue\n\n" + ("<<data>>\n" * 5))
    doc.add_table(rows=1, cols=2).rows[0].cells[0].text = "\n\nDean"
    doc.tables[-1].rows[0].cells[1].text = "\n\nMentor"
    doc.tables[-1].rows[0].cells[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    doc.add_paragraph("\ncc:\n-Notice Board\n-Departmental File\n-Mentoring File")
    doc.add_page_break()

    # Page 3: Action Taken
    doc.add_paragraph("\nDepartment of <<department>>\nCUIET – Applied Engineering\n\nAction Taken Report")
    doc.add_paragraph(
        "A Departmental Meeting was held to discuss the provision made and to formulate the adapted teaching methodology "
        "for the slow and advanced learners of <<semester>> (Batch: <<batch>>) students\n"
        "\nAction Taken for Slow Learners:\nExtra Classes\nList of Subjects to be offered:\nSubject Name (CODE)\n"
        "Subject Name (CODE)\n\nNote: The Classes will be offered as per the Schedule given in reference no. <<notice12>>.\n"
        "\nTeaching Methodology:\nLectures supported with visual modes\nAssignments, viva-voce, quizzes"
    )
    doc.add_paragraph("\nAction taken for Advanced Learners:\nNPTEL Lectures\nSkill enrichment")

    doc.add_paragraph("\n\nDate: <<date st1 next>>\nProgram Incharge\nDepartment of <<department>>")
    doc.add_page_break()

    # Page 4: Attendance
    doc.add_paragraph("\nDepartment of <<department>>\nCUIET – Applied Engineering\n\nAttendance")
    doc.add_paragraph("Subject Name: <<subject name>>\nSubject Code: <<subject code>>\nBatch/Semester:")
    doc.add_paragraph("\nS.No\tUID.no\tName\tDate\n" + "<<rows>>\n" * 5)
    doc.add_paragraph("\nSignature of Subject Incharge\nDepartment of <<department>>")
    doc.add_page_break()

    # Page 5: Performance
    doc.add_paragraph("\nDepartment of <<department>>\nCUIET – Applied Engineering\n\nPerformance Report")
    doc.add_paragraph("Slow Learners:\nClasses were provided and improvements are noted in ST-2.")
    doc.add_paragraph("Name\tUID\tSemester\tSubject Code\tST-1\t%\tST-2\t%")
    doc.add_paragraph("<<data>>\n" * 3)
    doc.add_paragraph("Advanced Learners: NPTEL courses were suggested for skill enhancement.")
    doc.add_paragraph("Date:\nProgram Incharge\nDepartment of <<department>>")

    doc.add_page_break()

    # Final ST-I + ST-II Based Notice
    doc.add_paragraph("Ref. No.: - CUIET/MED/SAL/<<year>>/<<semester>>/<<notice21>>\t\tDate: ")
    doc.add_paragraph("\nDepartment of <<department>>\nCUIET – Applied Engineering\n\nNotice")
    doc.add_paragraph(
        "\nSubject: Identification of Slow and Advanced Learners\n\n"
        "The students of <<Branch>> <<session>> were classified based on observations and performance in ST-I and ST-II. "
        "Below <<slow threshold>> in ST-2 = Slow, above <<advanced threshold>> = Advanced. See <<Annexure A2>>.\n\n"
        "Note: Mentors are requested to inform the above students."
    )
    doc.add_table(rows=1, cols=2).rows[0].cells[0].text = "\n\nDean"
    doc.tables[-1].rows[0].cells[1].text = "\n\nMentor"
    doc.tables[-1].rows[0].cells[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    doc.add_paragraph("\ncc:\n-Notice Board\n-Departmental File\n-Mentoring File")

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# === Streamlit App ===
st.title("Minimal Placeholder Notice Generator")

with st.sidebar:
    st.header("Inputs")
    year = st.text_input("Year", "2025")
    semester = st.text_input("Semester", "03")
    date = st.text_input("Date", "10-04-2025")
    logo = st.file_uploader("Upload Logo (optional)", type=["png", "jpg"])

data = {
    "year": year,
    "semester": semester,
    "date": date,
}

if st.button("Generate Placeholder Document"):
    file = generate_placeholder_doc(data, logo)
    st.success("Placeholder document ready!")
    st.download_button("📄 Download DOCX", data=file, file_name="Notice_Placeholder.docx")