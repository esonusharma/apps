# streamlit_app.py

import streamlit as st
import pandas as pd
from docx import Document
from io import BytesIO

def classify_student(percentage):
    if percentage < 40:
        return 'Slow Learner'
    elif percentage > 75:
        return 'Advanced Learner'
    else:
        return 'Average Learner'

def process_file(uploaded_file):
    df = pd.read_excel(uploaded_file)

    required_columns = ['Roll No.', 'Student Name', 'Subject Code', 'Branch', 'Batch', 'ST1 Percentage', 'ST2 Percentage']
    df = df[required_columns]

    df['ST1 Status'] = df['ST1 Percentage'].apply(classify_student)
    df['ST2 Status'] = df['ST2 Percentage'].apply(classify_student)

    return df

def create_word_tables(df, test_col, status_col, test_name):
    grouped = df.groupby(['Batch', 'Branch'])
    output = BytesIO()
    doc = Document()
    doc.add_heading(f'{test_name} Learner Classification', 0)

    for (batch, branch), group in grouped:
        doc.add_heading(f'Batch: {batch}, Branch: {branch}', level=1)
        table = doc.add_table(rows=1, cols=7)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Roll No.'
        hdr_cells[1].text = 'Student Name'
        hdr_cells[2].text = 'Subject Code'
        hdr_cells[3].text = 'Branch'
        hdr_cells[4].text = f'{test_name} Percentage'
        hdr_cells[5].text = 'Status'
        hdr_cells[6].text = 'Batch'

        for _, row in group.iterrows():
            row_cells = table.add_row().cells
            row_cells[0].text = str(row['Roll No.'])
            row_cells[1].text = str(row['Student Name'])
            row_cells[2].text = str(row['Subject Code'])
            row_cells[3].text = str(row['Branch'])
            row_cells[4].text = str(row[test_col])
            row_cells[5].text = str(row[status_col])
            row_cells[6].text = str(row['Batch'])

        doc.add_paragraph()  # Add space after each table

    doc.save(output)
    output.seek(0)
    return output

# Streamlit UI
st.title("Student Classification Report Generator")

uploaded_files = st.file_uploader("Upload Excel files", type=["xlsx"], accept_multiple_files=True)

if uploaded_files:
    all_data = pd.DataFrame()

    for file in uploaded_files:
        df = process_file(file)
        all_data = pd.concat([all_data, df], ignore_index=True)

    st.success("Files processed successfully!")

    # ST1 Report
    docx_st1 = create_word_tables(all_data, 'ST1 Percentage', 'ST1 Status', 'ST1')
    st.download_button(label="Download ST1 Learner Report", data=docx_st1, file_name="ST1_Report.docx")

    # ST2 Report
    docx_st2 = create_word_tables(all_data, 'ST2 Percentage', 'ST2 Status', 'ST2')
    st.download_button(label="Download ST2 Learner Report", data=docx_st2, file_name="ST2_Report.docx")