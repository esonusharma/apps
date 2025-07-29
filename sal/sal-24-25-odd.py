import streamlit as st
import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from io import BytesIO
import zipfile

# ========== SIDEBAR SETTINGS ==========
st.sidebar.header("📊 Learner Classification Settings")

slow_thresh = st.sidebar.number_input("Slow Learner Threshold (%)", min_value=0, max_value=100, value=40)
advanced_thresh = st.sidebar.number_input("Advanced Learner Threshold (%)", min_value=0, max_value=100, value=75)

uploaded_files = st.sidebar.file_uploader("📂 Upload Excel files", type=["xlsx"], accept_multiple_files=True)

# ========== FUNCTIONS ==========

def classify_student(percentage):
    if percentage < slow_thresh:
        return 'Slow Learner'
    elif percentage > advanced_thresh:
        return 'Advanced Learner'
    else:
        return 'Average Learner'



def sanitize_percent(value):
    if pd.isna(value):
        return 'A'
    str_val = str(value).strip().lower()
    if str_val in ['a', 'ab', 'absent']:
        return 'A'
    try:
        return float(value)
    except:
        return 'A'

def process_file(uploaded_file):
    df = pd.read_excel(uploaded_file)
    required_columns = [
        'Roll No.', 'Student Name', 'Subject Code', 'Subject Name',
        'Branch', 'Batch', 'Semester', 'ST1 Percentage', 'ST2 Percentage'
    ]
    df = df[required_columns]
    df['ST1 Percentage'] = df['ST1 Percentage'].apply(sanitize_percent)
    df['ST2 Percentage'] = df['ST2 Percentage'].apply(sanitize_percent)

    def classify_wrapper(val):
        if val == 'A':
            return 'Slow Learner'
        return classify_student(val)

    df['ST1 Status'] = df['ST1 Percentage'].apply(classify_wrapper)
    df['ST2 Status'] = df['ST2 Percentage'].apply(classify_wrapper)
    return df

def set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.append(tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for edge in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        elem = OxmlElement(f'w:{edge}')
        elem.set(qn('w:val'), 'single')
        elem.set(qn('w:sz'), '4')
        elem.set(qn('w:space'), '0')
        elem.set(qn('w:color'), '000000')
        tblBorders.append(elem)
    tblPr.append(tblBorders)

def add_heading(doc, text, size=12, align='center'):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.name = 'Times New Roman'
    if align == 'center':
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT

def format_cell_value(val):
    if pd.isna(val):
        return 'A'
    str_val = str(val).strip().lower()
    if str_val in ['a', 'ab', 'absent']:
        return 'A'
    try:
        num = float(val)
        if num.is_integer():
            return str(int(num))
        return f"{num:.2f}"
    except:
        return str(val)

def add_table(doc, df, columns):
    table = doc.add_table(rows=1, cols=len(columns))
    set_table_borders(table)
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(columns):
        hdr_cells[i].text = col
        for para in hdr_cells[i].paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        hdr_cells[i].vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for i, col in enumerate(columns):
            raw_val = row.get(col, '')
            value = format_cell_value(raw_val)
            cells[i].text = value
            for para in cells[i].paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cells[i].vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER
    return table

def annexure_A(doc, df, test_name, annexure_code):
    doc.add_page_break()
    add_heading(doc, f"Annexure {annexure_code}", 14)
    add_heading(doc, "Slow and Advanced Learners", 12)
    sample = df.iloc[0]
    add_heading(doc, f"List of Slow and Advanced Learners in {test_name} (Batch: {int(sample['Batch']) if isinstance(sample['Batch'], float) and sample['Batch'].is_integer() else sample['Batch']}, Semester: {int(sample['Semester']) if isinstance(sample['Semester'], float) and sample['Semester'].is_integer() else sample['Semester']})", 12, 'left')
    filtered = df[df[f'{test_name} Status'].isin(['Slow Learner', 'Advanced Learner']) & df['Branch'].isin(['AE', 'ME', 'ME Minor CSE'])]
    columns = ['Roll No.', 'Student Name', 'Subject Code', 'Subject Name', 'Branch', f'{test_name} Percentage', f'{test_name} Status']
    add_table(doc, filtered, columns)

def annexure_B(doc, df, test_name, annexure_code):
    doc.add_page_break()
    add_heading(doc, f"Annexure {annexure_code}", 14)
    add_heading(doc, "Time Table", 12)
    slow_df = df[df[f'{test_name} Status'] == 'Slow Learner']
    subjects = slow_df[['Subject Code', 'Subject Name', 'Semester', 'Branch']].drop_duplicates()
    subjects['Date'] = ''
    subjects['Time'] = '1000 HRS to 1600 HRS'
    subjects['Venue'] = 'online'
    columns = ['Subject Code', 'Subject Name', 'Semester', 'Date', 'Time', 'Venue']
    add_table(doc, subjects, columns)

def annexure_C(doc, df, test_name, annexure_code, other_test=None):
    doc.add_page_break()
    add_heading(doc, f"Annexure {annexure_code}", 14)
    add_heading(doc, "Attendance", 12)
    slow_df = df[df[f'{test_name} Status'] == 'Slow Learner']
    if other_test:
        slow_df['Attendance'] = slow_df.apply(lambda x: 'A' if x[f'{other_test} Status'] == 'Slow Learner' else 'P', axis=1)
    else:
        slow_df['Attendance'] = ''
    columns = ['Roll No.', 'Student Name', 'Subject Code', 'Subject Name', 'Branch', 'Attendance']
    add_table(doc, slow_df, columns)

def annexure_D(doc, df, test_name, annexure_code):
    doc.add_page_break()
    add_heading(doc, f"Annexure {annexure_code}", 14)
    add_heading(doc, "Performance", 12)
    if annexure_code.endswith("1"):
        filt = (df['ST1 Status'] == 'Slow Learner') & (df['ST2 Status'] != 'Slow Learner')
        temp = df[filt].copy()
        columns = ['Roll No.', 'Student Name', 'Subject Code', 'Subject Name', 'Branch', 'ST1 Percentage', 'ST2 Percentage']
    else:
        filt = (df['ST1 Status'] == 'Slow Learner') & (df['ST2 Status'] == 'Slow Learner')
        temp = df[filt].copy()
        temp['ETE Grade'] = ''
        columns = ['Roll No.', 'Student Name', 'Subject Code', 'Subject Name', 'Branch', 'ETE Grade']
    add_table(doc, temp, columns)

def generate_grouped_docs(df):
    zip_buffer = BytesIO()
    with zipfile.ZipFile(zip_buffer, "w") as zipf:
        grouped = df.groupby(['Batch', 'Branch', 'Semester'])
        for (batch, branch, semester), group in grouped:
            if branch not in ['AE', 'ME', 'ME Minor CSE']:
                continue
            doc = Document()
            doc.styles['Normal'].font.name = 'Times New Roman'
            doc.styles['Normal'].font.size = Pt(12)

            annexure_A(doc, group, 'ST1', 'A1')
            annexure_B(doc, group, 'ST1', 'B1')
            annexure_C(doc, group, 'ST1', 'C1', 'ST2')
            annexure_D(doc, group, 'ST1', 'D1')
            annexure_A(doc, group, 'ST2', 'A2')
            annexure_B(doc, group, 'ST2', 'B2')
            annexure_C(doc, group, 'ST2', 'C2')
            annexure_D(doc, group, 'ST2', 'D2')

            doc_buffer = BytesIO()
            doc.save(doc_buffer)
            doc_buffer.seek(0)

            session = str(int(group['Batch'].iloc[0])) if isinstance(group['Batch'].iloc[0], float) and group['Batch'].iloc[0].is_integer() else str(group['Batch'].iloc[0])
            filename = f"{session}_{branch}_{int(semester) if isinstance(semester, float) and semester.is_integer() else semester}_Annexures.docx"
            zipf.writestr(filename, doc_buffer.read())

    zip_buffer.seek(0)
    return zip_buffer

# ========== MAIN DISPLAY AREA ==========

st.title("📚 Annexure Generator for Slow and Advanced Learners")

if uploaded_files:
    all_data = pd.DataFrame()
    for file in uploaded_files:
        df = process_file(file)
        all_data = pd.concat([all_data, df], ignore_index=True)
    st.success("✅ Files processed successfully!")

    final_zip = generate_grouped_docs(all_data)

    st.download_button(
        label="📦 Download All Annexures as ZIP",
        data=final_zip,
        file_name="All_Annexures.zip",
        mime="application/zip"
    )
