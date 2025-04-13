# learner_report_app.py

import streamlit as st
import pandas as pd
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from io import BytesIO
import zipfile
from datetime import date

# ====== UI: Sidebar Inputs ======
st.sidebar.title("📄 Report Settings")

selected_year = st.sidebar.selectbox("Select Year", ["2023", "2024", "2025"])
department = st.sidebar.selectbox("Department", [
    "Mechanical Engineering", "Mechatronics Engineering", "Civil Engineering",
    "Electrical Engineering", "Computer Science and Engineering"
])

notice11_date = st.sidebar.date_input("Date for Notice 11", date.today())
notice12_date = st.sidebar.date_input("Date for Notice 12", date.today())
notice21_date = st.sidebar.date_input("Date for Notice 21", date.today())
notice22_date = st.sidebar.date_input("Date for Notice 22", date.today())
date_st1_next = st.sidebar.date_input("Action Taken Date (ST1)", date.today())

slow_thresh = st.sidebar.slider("Slow Learner Threshold (%)", 0, 100, 40)
adv_thresh = st.sidebar.slider("Advanced Learner Threshold (%)", 0, 100, 75)

uploaded_files = st.sidebar.file_uploader("📂 Upload Excel files", type=["xlsx"], accept_multiple_files=True)

# ====== Helper Functions ======
def sanitize_percent(val):
    try:
        return float(str(val).strip().lower().replace('%', ''))
    except:
        return 0.0

def classify(p):
    if p < slow_thresh:
        return "Slow Learner"
    elif p > adv_thresh:
        return "Advanced Learner"
    return "Average Learner"

def set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr or OxmlElement('w:tblPr')
    tblBorders = OxmlElement('w:tblBorders')
    for edge in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        elem = OxmlElement(f'w:{edge}')
        elem.set(qn('w:val'), 'single')
        elem.set(qn('w:sz'), '4')
        elem.set(qn('w:space'), '0')
        elem.set(qn('w:color'), '000000')
        tblBorders.append(elem)
    tblPr.append(tblBorders)
    tbl.tblPr = tblPr

def write_paragraph(doc, text, align='left', bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.bold = bold
    if align == 'center':
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    elif align == 'right':
        p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    return p

def write_notice_header(doc, ref_no, date_text):
    table = doc.add_table(rows=1, cols=2)
    table.autofit = True
    row = table.rows[0].cells
    row[0].text = ref_no
    row[1].text = f"Date: {date_text}"
    row[0].paragraphs[0].runs[0].font.name = 'Times New Roman'
    row[0].paragraphs[0].runs[0].font.size = Pt(12)
    row[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    row[1].paragraphs[0].runs[0].font.name = 'Times New Roman'
    row[1].paragraphs[0].runs[0].font.size = Pt(12)

def add_annexure_a1(doc, df):
    doc.add_page_break()
    write_paragraph(doc, "Annexure A1", align='center', bold=True)
    filtered = df[df['ST1 Status'].isin(['Slow Learner', 'Advanced Learner'])]
    if filtered.empty:
        write_paragraph(doc, "No slow or advanced learners identified in ST-I.")
        return

    table = doc.add_table(rows=1, cols=8)
    table.style = 'Table Grid'
    set_table_borders(table)
    hdr = ["Roll No.", "Student Name", "Subject Code", "Subject Name", "Branch", "ST1 %", "Status", "Batch"]
    for i, h in enumerate(hdr):
        table.cell(0, i).text = h
    for _, row in filtered.iterrows():
        cells = table.add_row().cells
        cells[0].text = str(row['Roll No.'])
        cells[1].text = str(row['Student Name'])
        cells[2].text = str(row['Subject Code'])
        cells[3].text = str(row['Subject Name'])
        cells[4].text = str(row['Branch'])
        cells[5].text = f"{row['ST1 Percentage']:.1f}"
        cells[6].text = str(row['ST1 Status'])
        cells[7].text = str(row['Batch'])

# ====== Main Document Generator ======
def generate_document_for_group(group, meta):
    branch = group['Branch'].iloc[0]
    batch = group['Batch'].iloc[0]
    session = f"{batch}–{int(batch)+4}"
    semester = '3rd'  # for now, assume 3rd semester; can extract if available

    doc = Document()
    doc.styles['Normal'].font.name = 'Times New Roman'
    doc.styles['Normal'].font.size = Pt(12)

    # PAGE 1: ST-I Identification Notice
    write_notice_header(doc,
        f"Ref. No.: CUIET/MED/SAL/{meta['year']}/{semester}/{branch}/01",
        meta['notice11_date'])
    write_paragraph(doc, f"Department of {meta['department']}", align='center')
    write_paragraph(doc, "CUIET – Applied Engineering", align='center')
    write_paragraph(doc, "\nNotice", align='center', bold=True)
    write_paragraph(doc, "\nSubject: Identification of Slow and Advanced Learners", bold=True)

    write_paragraph(doc,
        f"The below mentioned students of {branch} {session} were classified into the Slow and Advanced learners' categories "
        "based upon the observations and feedback from the mentors, teachers and academic performance in ST-I. "
        f"Students, who score marks below {slow_thresh}% are categorized as slow learners and above {adv_thresh}% are categorized as advanced learners. "
        "These distinguished parameters enabled in identification of advanced learners and slow learners. "
        "The details of slow and advanced learners is available in Annexure A1.")

    write_paragraph(doc, "Note: Mentors are requested to inform the above students.")
    write_paragraph(doc, "\n\nDean", align='left')
    write_paragraph(doc, "\n\nMentor", align='right')
    write_paragraph(doc, "\ncc:\n-Notice Board\n-Departmental File\n-Mentoring File")

    # PAGE 2: Annexure A1
    add_annexure_a1(doc, group)

    return doc

# ====== Streamlit UI + Generation ======
st.title("📘 Learner Classification Document Generator")

if st.button("Generate Reports"):
    if not uploaded_files:
        st.warning("Please upload at least one Excel file.")
    else:
        zip_buffer = BytesIO()
        with zipfile.ZipFile(zip_buffer, 'w') as zipf:
            for f in uploaded_files:
                df = pd.read_excel(f)
                df['ST1 Percentage'] = df['ST1 Percentage'].apply(sanitize_percent)
                df['ST2 Percentage'] = df['ST2 Percentage'].apply(sanitize_percent)
                df['ST1 Status'] = df['ST1 Percentage'].apply(classify)
                df['ST2 Status'] = df['ST2 Percentage'].apply(classify)

                grouped = df.groupby(['Batch', 'Branch'])
                for (batch, branch), group in grouped:
                    doc = generate_document_for_group(group, {
                        'year': selected_year,
                        'department': department,
                        'notice11_date': notice11_date.strftime('%d-%m-%Y')
                    })
                    buffer = BytesIO()
                    doc.save(buffer)
                    filename = f"{batch}_{branch}.docx"
                    zipf.writestr(filename, buffer.getvalue())

        zip_buffer.seek(0)
        st.download_button("📥 Download All Reports", data=zip_buffer, file_name="Learner_Reports.zip")
